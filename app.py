from flask import Flask, request, send_file, render_template_string
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import Image as RLImage
from docx import Document as DocxDocument
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, re, os

app = Flask(__name__)

LOGO1 = os.path.join(os.path.dirname(__file__), 'logo1.png')
LOGO2 = os.path.join(os.path.dirname(__file__), 'logo2.png')

# ─── Texto das atas ────────────────────────────────────────────────────────────

def texto_defesa(d):
    nivel = 'Mestre' if d['nivel'] == 'mestrado' else 'Doutor'
    trab  = 'dissertação' if d['nivel'] == 'mestrado' else 'tese'
    return [(
        f"Aos {d['dia']} dias do mês de {d['mes']}, do ano de dois mil e vinte e {d['ano']}, "
        f"às {d['hora']}, {d['modalidade']}, realizou-se a Defesa do(a) discente "
        f"{d['discente']}, como parte das exigências para obtenção do título de {nivel}, com a "
        f"{trab} intitulada: \u201c{d['titulo']}\u201d do Curso de Pós-graduação Stricto Sensu em "
        f"Ciências Ambientais, perante a banca examinadora, composta pelos(as) examinadores(as) "
        f"listados(as) abaixo, onde a sessão foi aberta pelo(a) presidente "
        f"{d['or_trat']} {d['or_nome']} e, após apresentação, o(a) discente foi arguido(a) "
        f"pela Banca Examinadora. Em sessão secreta foi decidido o resultado da defesa, sendo "
        f"o(a) discente considerado(a) {d['resultado']}. Encerrada a sessão secreta, o Presidente "
        f"informou o resultado. Nada mais havendo a tratar, eu, Presidente da banca, lavrei a "
        f"presente ata que assino juntamente com os membros da Banca Examinadora. Para a obtenção "
        f"do título ainda é necessário o cumprimento das exigências contidas no Regimento do "
        f"Programa de Pós-graduação Stricto Sensu em Ciências Ambientais - Resolução n.º {d['resolucao']}."
    )]

def texto_qualificacao(d):
    nivel = 'Mestre' if d['nivel'] == 'mestrado' else 'doutor'
    trab  = 'dissertação' if d['nivel'] == 'mestrado' else 'tese'
    return [
        (f"Aos {d['dia']} dias do mês de {d['mes']} do ano de dois mil e vinte e {d['ano']}, "
         f"às {d['hora']}, {d['modalidade']} realizou-se o Exame de Qualificação do(a) discente "
         f"{d['discente']}, como parte das exigências para obtenção do título de {nivel}, com a "
         f"versão preliminar da {trab} intitulada: \u201c{d['titulo']}\u201d do Curso de "
         f"Pós-graduação Stricto Sensu em Ciências Ambientais, perante a banca examinadora, "
         f"composta pelos professores abaixo."),
        '',
        (f"Após apresentação e arguição, a banca examinadora conclui pela APROVAÇÃO do(a) "
         f"discente com conceito final do Exame de Qualificação: {d['conceito']}"),
        '',
        'I. \u201cA\u201d \u2013 aprovação, considerando pequenas reformulações sugeridas pela banca;',
        'II. \u201cB\u201d \u2013 aprovação, com reformulações estruturais de acordo com as sugestões da Banca;',
        'III. \u201cC\u201d \u2013 aprovação, com reformulações estruturais e metodológicas apresentadas pela Banca;',
        'IV. \u201cD\u201d \u2013 Reprovação e recomendação de ampla reformulação para novo Exame de Qualificação.',
        '',
        ('Em seguida a presidente da banca agradeceu a participação dos presentes e deu por '
         'encerrada a presente reunião, a tudo presenciei, lavrei e assinei a presente ata.'),
    ]

TITLES = {
    ('defesa','mestrado'):        'ATA DE DEFESA DE DISSERTAÇÃO',
    ('defesa','doutorado'):       'ATA DE DEFESA DE TESE',
    ('qualificacao','mestrado'):  'ATA DO EXAME DE QUALIFICAÇÃO DE DISSERTAÇÃO',
    ('qualificacao','doutorado'): 'ATA DE EXAME DE QUALIFICAÇÃO DE DOUTORADO',
}

# ─── Gerar PDF ─────────────────────────────────────────────────────────────────

def gerar_pdf(d, membros, titulo, linhas):
    buf = io.BytesIO()
    PAGE_W, PAGE_H = A4
    ML, MR, MT, MB = 25*mm, 20*mm, 18*mm, 20*mm
    W = PAGE_W - ML - MR

    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=ML, rightMargin=MR, topMargin=MT, bottomMargin=MB)

    s_ib  = ParagraphStyle('ib', fontName='Helvetica-Bold', fontSize=9,  alignment=TA_CENTER, leading=12)
    s_i   = ParagraphStyle('i',  fontName='Helvetica',      fontSize=8,  alignment=TA_CENTER, leading=11)
    s_tit = ParagraphStyle('t',  fontName='Times-Bold',     fontSize=13, alignment=TA_CENTER, leading=17, spaceBefore=8, spaceAfter=12)
    s_b   = ParagraphStyle('b',  fontName='Times-Roman',    fontSize=11, alignment=TA_JUSTIFY, leading=17, firstLineIndent=18, spaceAfter=6)
    s_l   = ParagraphStyle('l',  fontName='Times-Roman',    fontSize=11, alignment=TA_JUSTIFY, leading=16, spaceAfter=4)
    s_sn  = ParagraphStyle('sn', fontName='Times-Bold',     fontSize=9,  alignment=TA_CENTER, leading=12)
    s_si  = ParagraphStyle('si', fontName='Times-Roman',    fontSize=8,  alignment=TA_CENTER, leading=11, textColor=colors.Color(.3,.3,.3))
    s_sc  = ParagraphStyle('sc', fontName='Times-Roman',    fontSize=7.5,alignment=TA_CENTER, leading=10, textColor=colors.Color(.4,.4,.4))
    s_ft  = ParagraphStyle('ft', fontName='Helvetica',      fontSize=7,  alignment=TA_CENTER, leading=9,  textColor=colors.Color(.6,.6,.6))

    story = []

    # Header
    logo1_img = RLImage(LOGO1, width=18*mm, height=20*mm)
    logo2_img = RLImage(LOGO2, width=18*mm, height=20*mm)
    htbl = Table([[logo1_img, [
        Paragraph('UNIVERSIDADE DO ESTADO DE MATO GROSSO', s_ib),
        Paragraph('Programa de Pós-Graduação Stricto Sensu em Ciências Ambientais', s_i),
        Paragraph('Campus de Cáceres – Alta Floresta – Nova Xavantina', s_i),
    ], logo2_img]], colWidths=[22*mm, W-44*mm, 22*mm])
    htbl.setStyle(TableStyle([
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('LEFTPADDING',(0,0),(-1,-1),0),('RIGHTPADDING',(0,0),(-1,-1),0),
        ('TOPPADDING',(0,0),(-1,-1),0),('BOTTOMPADDING',(0,0),(-1,-1),0),
    ]))
    story.append(htbl)
    story.append(HRFlowable(width=W, thickness=0.5, color=colors.Color(.7,.7,.7), spaceAfter=4))
    story.append(Paragraph(titulo, s_tit))

    # Corpo
    for linha in linhas:
        if linha == '':
            story.append(Spacer(1, 5))
        elif any(linha.startswith(x) for x in ['I.','II.','III.','IV.']):
            story.append(Paragraph(linha, s_l))
        else:
            story.append(Paragraph(linha, s_b))

    story.append(Spacer(1, 14))

    # Assinaturas
    SIG_H = 42*mm
    def sig_cell(m):
        return [
            Spacer(1, SIG_H - 20*mm),
            Paragraph('_' * 44, s_sn),
            Paragraph(f'{m["trat"]} {m["nome"]}', s_sn),
            Paragraph(m['inst'], s_si),
            (Paragraph(f'CPF: {m["cpf"]}', s_sc) if m.get('cpf') else Spacer(1,4)),
        ]

    rows, num_rows = [], (len(membros)+1)//2
    for i in range(0, len(membros), 2):
        a, b = membros[i], (membros[i+1] if i+1 < len(membros) else None)
        rows.append([sig_cell(a), sig_cell(b) if b else [Spacer(1, SIG_H)]])

    stbl = Table(rows, colWidths=[W/2-3*mm, W/2-3*mm], rowHeights=[SIG_H]*len(rows))
    stbl.setStyle(TableStyle([
        ('BOX',(0,0),(-1,-1),0.5,colors.Color(.6,.6,.6)),
        ('INNERGRID',(0,0),(-1,-1),0.5,colors.Color(.6,.6,.6)),
        ('VALIGN',(0,0),(-1,-1),'BOTTOM'),('ALIGN',(0,0),(-1,-1),'CENTER'),
        ('LEFTPADDING',(0,0),(-1,-1),6),('RIGHTPADDING',(0,0),(-1,-1),6),
        ('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),6),
        ('BACKGROUND',(0,0),(-1,-1),colors.white),
    ]))
    story.append(stbl)
    story.append(Spacer(1,8))
    story.append(Paragraph(
        'Concepção: Prof. Dr. Ernandes Sobreira Oliveira Junior · Biólogo · UNEMAT — Plataforma de Atas PPGCA', s_ft))

    doc.build(story)
    buf.seek(0)
    return buf

# ─── Gerar DOCX ────────────────────────────────────────────────────────────────

def gerar_docx(d, membros, titulo, linhas):
    buf = io.BytesIO()
    doc = DocxDocument()
    sec = doc.sections[0]
    sec.page_width=Mm(210); sec.page_height=Mm(297)
    sec.left_margin=Mm(25); sec.right_margin=Mm(20)
    sec.top_margin=Mm(18);  sec.bottom_margin=Mm(20)

    def add_p(text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, font='Arial',
              space_before=0, space_after=4, first_indent=None, italic=False, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if first_indent is not None:
            p.paragraph_format.first_line_indent = Mm(first_indent)
        run = p.add_run(text)
        run.bold=bold; run.italic=italic
        run.font.size=Pt(size); run.font.name=font
        if color: run.font.color.rgb = RGBColor(*color)
        return p

    add_p('UNIVERSIDADE DO ESTADO DE MATO GROSSO', bold=True, size=10)
    add_p('Programa de Pós-Graduação Stricto Sensu em Ciências Ambientais', size=9)
    add_p('Campus de Cáceres – Alta Floresta – Nova Xavantina', size=9, space_after=6)

    # linha separadora
    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'AAAAAA')
    pBdr.append(bot); pPr.append(pBdr)

    add_p(titulo, bold=True, size=13, font='Times New Roman', space_before=8, space_after=12)

    for linha in linhas:
        if linha == '':
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
            continue
        is_item = any(linha.startswith(x) for x in ['I.','II.','III.','IV.'])
        add_p(linha, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              font='Times New Roman', space_after=5,
              first_indent=(None if is_item else 10))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Tabela de assinaturas
    n_rows = (len(membros)+1)//2
    table = doc.add_table(rows=n_rows, cols=2)
    table.style = 'Table Grid'

    for row in table.rows:
        row.height = Mm(42)
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            for side in ['top','left','bottom','right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
                b.set(qn('w:space'),'0');    b.set(qn('w:color'),'AAAAAA')
                tcPr.append(b)

    for idx, m in enumerate(membros):
        ri, ci = idx//2, idx%2
        cell = table.cell(ri, ci)
        for para in cell.paragraphs:
            for run in para.runs: run.text = ''

        def cp(text='', bold=False, size=9, color=None):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if text:
                run = p.add_run(text)
                run.bold=bold; run.font.size=Pt(size); run.font.name='Times New Roman'
                if color: run.font.color.rgb=RGBColor(*color)
            return p

        for _ in range(3): cp()           # espaço para assinatura
        cp('_'*42)                         # linha
        cp(f'{m["trat"]} {m["nome"]}', bold=True)
        cp(m['inst'], color=(80,80,80))
        if m.get('cpf'): cp(f'CPF: {m["cpf"]}', size=8, color=(100,100,100))

    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    add_p('Concepção: Prof. Dr. Ernandes Sobreira Oliveira Junior · Biólogo · UNEMAT',
          size=7, italic=True, color=(150,150,150))

    doc.save(buf)
    buf.seek(0)
    return buf

# ─── Parsear membros do formulário ─────────────────────────────────────────────

def parse_membros(form, prefix):
    membros = []
    # orientador primeiro
    membros.append({
        'trat': form.get(f'{prefix}_or_trat','Prof. Dr.'),
        'nome': form.get(f'{prefix}_or_nome',''),
        'cpf':  form.get(f'{prefix}_or_cpf',''),
        'inst': form.get(f'{prefix}_or_inst','UNEMAT'),
    })
    i = 1
    while True:
        nome = form.get(f'{prefix}_m{i}_nome','').strip()
        if not nome: break
        membros.append({
            'trat': form.get(f'{prefix}_m{i}_trat','Prof. Dr.'),
            'nome': nome,
            'cpf':  form.get(f'{prefix}_m{i}_cpf',''),
            'inst': form.get(f'{prefix}_m{i}_inst','UNEMAT'),
        })
        i += 1
    return membros

# ─── HTML da plataforma ────────────────────────────────────────────────────────

HTML = '''<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Gerador de Atas – PPGCA/UNEMAT</title>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&family=DM+Sans:wght@400;500;600&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'DM Sans',sans-serif;background:linear-gradient(135deg,#0e3d26,#1a5c3a 50%,#1d6640);min-height:100vh}
.hero{padding:28px 20px 0;text-align:center}
.hero-logos{display:flex;justify-content:center;align-items:center;gap:20px;margin-bottom:14px;flex-wrap:wrap}
.hero-logos img{height:66px;object-fit:contain;filter:drop-shadow(0 2px 8px rgba(0,0,0,.4))}
.hero h1{font-family:'Playfair Display',serif;font-size:clamp(1.3rem,4vw,2.1rem);color:#fff;margin-bottom:5px}
.hero p{font-size:.86rem;color:rgba(255,255,255,.75);margin-bottom:3px}
.credit{font-size:.75rem;color:#f0d98a;margin-bottom:20px;font-style:italic}
.tab-bar{display:flex;justify-content:center;gap:4px;flex-wrap:wrap;padding:0 12px}
.tab-btn{background:rgba(255,255,255,.12);color:rgba(255,255,255,.82);border:1.5px solid rgba(255,255,255,.18);border-bottom:none;border-radius:10px 10px 0 0;padding:9px 14px;font-family:'DM Sans',sans-serif;font-size:.81rem;font-weight:500;cursor:pointer;transition:.2s;white-space:nowrap}
.tab-btn:hover{background:rgba(255,255,255,.2);color:#fff}
.tab-btn.active{background:#fff;color:#1a5c3a;font-weight:700;border-color:#fff}
.card{background:#fff;border-radius:0 0 12px 12px;box-shadow:0 8px 40px rgba(26,92,58,.18);max-width:820px;margin:0 auto;padding:26px 30px 30px;position:relative}
.card::before{content:'';position:absolute;top:0;left:30px;right:30px;height:3px;background:linear-gradient(90deg,#1a5c3a,#c8a84b,#1a5c3a);border-radius:0 0 4px 4px}
.wrap{max-width:820px;margin:0 auto;padding:0 12px 48px}
.sec{font-family:'Playfair Display',serif;font-size:.95rem;color:#1a5c3a;margin:18px 0 9px;padding-bottom:5px;border-bottom:1.5px solid #e8f4ee;display:flex;align-items:center;gap:7px}
.sec::before{content:'';width:4px;height:14px;background:#c8a84b;border-radius:2px;flex-shrink:0}
.row{display:grid;gap:11px}
.r2{grid-template-columns:1fr 1fr}
.r3{grid-template-columns:1fr 1fr 1fr}
.f{display:flex;flex-direction:column;gap:4px;margin-bottom:10px}
.f label{font-size:.7rem;font-weight:700;color:#374151;text-transform:uppercase;letter-spacing:.06em}
.f input,.f select{border:1.5px solid #d1d5db;border-radius:8px;padding:9px 11px;font-family:'DM Sans',sans-serif;font-size:.87rem;color:#111;background:#f9fafb;outline:none;width:100%;transition:.15s}
.f input:focus,.f select:focus{border-color:#1a5c3a;box-shadow:0 0 0 3px rgba(26,92,58,.1);background:#fff}
.mb{background:#f9fafb;border:1.5px solid #eee;border-radius:10px;padding:11px 13px 3px;margin-bottom:8px}
.mbl{font-size:.7rem;font-weight:700;color:#1a5c3a;text-transform:uppercase;letter-spacing:.08em;margin-bottom:8px;display:flex;align-items:center;gap:5px}
.badge{background:#1a5c3a;color:#fff;font-size:.62rem;font-weight:700;padding:2px 7px;border-radius:20px}
.btn-add{background:#fff;color:#1a5c3a;border:1.5px solid #1a5c3a;border-radius:8px;padding:7px 13px;font-family:'DM Sans',sans-serif;font-size:.79rem;font-weight:600;cursor:pointer;margin:3px 0 6px;transition:.15s}
.btn-add:hover{background:#e8f4ee}
.btn-row{display:flex;gap:8px;margin-top:22px;flex-wrap:wrap}
.btn{padding:13px 26px;border-radius:10px;font-family:'DM Sans',sans-serif;font-size:.92rem;font-weight:700;cursor:pointer;border:none;display:flex;align-items:center;gap:7px;transition:.18s;text-decoration:none;justify-content:center}
.btn-green{background:linear-gradient(135deg,#1a5c3a,#2d7a52);color:#fff;box-shadow:0 4px 16px rgba(26,92,58,.25)}
.btn-green:hover{transform:translateY(-1px);box-shadow:0 6px 24px rgba(26,92,58,.35)}
.btn-gold{background:linear-gradient(135deg,#c8a84b,#b8943b);color:#fff;box-shadow:0 4px 16px rgba(200,168,75,.3)}
.btn-gold:hover{transform:translateY(-1px)}
.tab-content{display:none}
.tab-content.active{display:block}
@media(max-width:580px){.card{padding:15px 12px}.r2,.r3{grid-template-columns:1fr}}
</style>
</head>
<body>
<div class="hero">
  <div class="hero-logos">
    <img src="/logo1" alt="MT">
    <img src="/logo2" alt="UNEMAT">
    <img src="/logo3" alt="CA">
  </div>
  <h1>Gerador de Atas Acadêmicas</h1>
  <p>Programa de Pós-Graduação Stricto Sensu em Ciências Ambientais – UNEMAT</p>
  <p class="credit">💡 Concepção: Prof. Dr. Ernandes Sobreira Oliveira Junior · Biólogo · UNEMAT &nbsp;|&nbsp; Moldado com IA</p>
</div>

<div class="wrap">
  <div class="tab-bar">
    <button class="tab-btn active" onclick="setTab('qm',this)">📋 Qualif. Mestrado</button>
    <button class="tab-btn" onclick="setTab('qd',this)">📋 Qualif. Doutorado</button>
    <button class="tab-btn" onclick="setTab('dm',this)">🎓 Defesa Mestrado</button>
    <button class="tab-btn" onclick="setTab('dd',this)">🎓 Defesa Doutorado</button>
  </div>

  <div class="card">

    <!-- ===== QUALIF MESTRADO ===== -->
    <form id="form-qm" action="/gerar" method="POST" target="_blank">
    <input type="hidden" name="tipo" value="qualificacao">
    <input type="hidden" name="nivel" value="mestrado">
    <div id="tab-qm" class="tab-content active">
      {{ form_base('qm', 'dissertação') }}
      {{ conceito_field('qm') }}
      <input type="hidden" name="qm_resolucao" value="10/2026-CONSUNI">
    </div>
    <div class="btn-row">
      <button type="submit" name="fmt" value="pdf" class="btn btn-green" form="form-qm">⬇ Baixar PDF</button>
      <button type="submit" name="fmt" value="docx" class="btn btn-gold" form="form-qm">📄 Baixar Word</button>
    </div>
    </form>

    <!-- ===== QUALIF DOUTORADO ===== -->
    <form id="form-qd" action="/gerar" method="POST" target="_blank">
    <input type="hidden" name="tipo" value="qualificacao">
    <input type="hidden" name="nivel" value="doutorado">
    <div id="tab-qd" class="tab-content">
      {{ form_base('qd', 'tese') }}
      {{ conceito_field('qd') }}
      <input type="hidden" name="qd_resolucao" value="046/2024-CONSUNI">
    </div>
    </form>

    <!-- ===== DEFESA MESTRADO ===== -->
    <form id="form-dm" action="/gerar" method="POST" target="_blank">
    <input type="hidden" name="tipo" value="defesa">
    <input type="hidden" name="nivel" value="mestrado">
    <div id="tab-dm" class="tab-content">
      {{ form_base('dm', 'dissertação') }}
      {{ resultado_field('dm') }}
      <input type="hidden" name="dm_resolucao" value="10/2026-CONSUNI">
    </div>
    </form>

    <!-- ===== DEFESA DOUTORADO ===== -->
    <form id="form-dd" action="/gerar" method="POST" target="_blank">
    <input type="hidden" name="tipo" value="defesa">
    <input type="hidden" name="nivel" value="doutorado">
    <div id="tab-dd" class="tab-content">
      {{ form_base('dd', 'tese') }}
      {{ resultado_field('dd') }}
      <input type="hidden" name="dd_resolucao" value="046/2024-CONSUNI">
    </div>
    </form>

  </div>
</div>

<script>
var activeTab = 'qm';
function setTab(id, btn) {
  document.querySelectorAll('.tab-content').forEach(function(e){ e.classList.remove('active'); });
  document.querySelectorAll('.tab-btn').forEach(function(b){ b.classList.remove('active'); });
  document.getElementById('tab-'+id).classList.add('active');
  btn.classList.add('active');
  activeTab = id;
  // Move btn-row into active form area
  var btnRow = document.querySelector('.btn-row');
  var activeForm = document.getElementById('form-'+id);
  activeForm.appendChild(btnRow);
  // Update btn form attributes
  document.querySelectorAll('.btn-row .btn').forEach(function(b){
    b.setAttribute('form','form-'+id);
  });
}
function addMember(prefix, n) {
  var container = document.getElementById('members-'+prefix);
  var idx = container.children.length + 1;
  var div = document.createElement('div');
  div.className = 'mb';
  div.innerHTML = '<div class="mbl"><span class="badge">'+(idx+1)+'</span> Membro '+idx
    +' <button type="button" onclick="this.closest(\'.mb\').remove()" style="margin-left:auto;background:none;border:none;cursor:pointer;color:#bbb;font-size:.9rem">✕</button></div>'
    +'<div class="row r3">'
    +'<div class="f"><label>Tratamento</label><select name="'+prefix+'_m'+idx+'_trat"><option>Prof. Dr.</option><option>Profa. Dra.</option><option>Prof. Me.</option><option>Profa. Me.</option></select></div>'
    +'<div class="f"><label>Nome</label><input type="text" name="'+prefix+'_m'+idx+'_nome" placeholder="Nome completo"></div>'
    +'<div class="f"><label>CPF</label><input type="text" name="'+prefix+'_m'+idx+'_cpf" placeholder="000.000.000-00"></div>'
    +'</div>'
    +'<div class="row r2">'
    +'<div class="f"><label>Instituição</label><input type="text" name="'+prefix+'_m'+idx+'_inst" value="UNEMAT"></div>'
    +'</div>';
  container.appendChild(div);
}
// Move btn-row to initial form on load
window.onload = function() {
  var btnRow = document.querySelector('.btn-row');
  document.getElementById('form-qm').appendChild(btnRow);
};
</script>
</body>
</html>'''

def make_member_block(prefix, n, label):
    return f'''
    <div class="mb" id="{prefix}-mb-{n}">
      <div class="mbl"><span class="badge">{n+1}</span> {label}</div>
      <div class="row r3">
        <div class="f"><label>Tratamento</label>
          <select name="{prefix}_m{n}_trat">
            <option>Prof. Dr.</option><option>Profa. Dra.</option>
            <option>Prof. Me.</option><option>Profa. Me.</option>
          </select></div>
        <div class="f"><label>Nome</label>
          <input type="text" name="{prefix}_m{n}_nome" placeholder="Nome completo"></div>
        <div class="f"><label>CPF</label>
          <input type="text" name="{prefix}_m{n}_cpf" placeholder="000.000.000-00"></div>
      </div>
      <div class="row r2">
        <div class="f"><label>Instituição</label>
          <input type="text" name="{prefix}_m{n}_inst" value="UNEMAT"></div>
      </div>
    </div>'''

def make_form(prefix, trab):
    members_default = {
        'qm': [('Membro 1',''),('Membro 2','')],
        'qd': [('Membro 1',''),('Membro 2',''),('Membro 3','')],
        'dm': [('Membro 1',''),('Membro 2',''),('Membro 3','')],
        'dd': [('Membro 1',''),('Membro 2',''),('Membro 3',''),('Membro 4','')],
    }
    mbs = ''.join(make_member_block(prefix, i+1, label) for i,(label,_) in enumerate(members_default[prefix]))
    return f'''
    <div class="sec">📅 Data e Hora</div>
    <div class="row r3">
      <div class="f"><label>Dia (por extenso) *</label>
        <input type="text" name="dia" required placeholder="vinte e cinco"></div>
      <div class="f"><label>Mês *</label>
        <select name="mes" required>
          <option value="janeiro">Janeiro</option><option value="fevereiro">Fevereiro</option>
          <option value="março">Março</option><option value="abril">Abril</option>
          <option value="maio">Maio</option><option value="junho">Junho</option>
          <option value="julho">Julho</option><option value="agosto">Agosto</option>
          <option value="setembro">Setembro</option><option value="outubro">Outubro</option>
          <option value="novembro">Novembro</option><option value="dezembro">Dezembro</option>
        </select></div>
      <div class="f"><label>Ano (ex: cinco) *</label>
        <input type="text" name="ano" required placeholder="cinco"></div>
    </div>
    <div class="row r2">
      <div class="f"><label>Horário (por extenso) *</label>
        <input type="text" name="hora" required placeholder="quatorze horas"></div>
      <div class="f"><label>Modalidade</label>
        <select name="modalidade">
          <option value="de forma virtual">Virtual (online)</option>
          <option value="presencialmente">Presencial</option>
          <option value="em formato híbrido">Híbrido</option>
        </select></div>
    </div>

    <div class="sec">👤 Discente</div>
    <div class="f"><label>Nome completo *</label>
      <input type="text" name="discente" required placeholder="Nome completo do(a) discente"></div>
    <div class="f"><label>Título da {trab} *</label>
      <input type="text" name="titulo" required placeholder="Título completo"></div>

    <div class="sec">👨‍🏫 Orientador(a) – Presidente da Banca</div>
    <div class="mb">
      <div class="mbl"><span class="badge">P</span> Orientador(a) / Presidente</div>
      <div class="row r3">
        <div class="f"><label>Tratamento</label>
          <select name="{prefix}_or_trat">
            <option>Prof. Dr.</option><option>Profa. Dra.</option>
            <option>Prof. Me.</option><option>Profa. Me.</option>
          </select></div>
        <div class="f"><label>Nome *</label>
          <input type="text" name="{prefix}_or_nome" required placeholder="Nome completo"></div>
        <div class="f"><label>CPF</label>
          <input type="text" name="{prefix}_or_cpf" placeholder="000.000.000-00"></div>
      </div>
      <div class="row r2">
        <div class="f"><label>Instituição</label>
          <input type="text" name="{prefix}_or_inst" value="UNEMAT"></div>
      </div>
    </div>

    <div class="sec">👥 Demais Membros da Banca</div>
    <div id="members-{prefix}">{mbs}</div>
    <button type="button" class="btn-add" onclick="addMember('{prefix}')">+ Adicionar membro</button>

    <div class="sec">✅ Resultado</div>'''

def make_conceito(prefix):
    return f'''
    <div class="row r2">
      <div class="f"><label>Conceito Final *</label>
        <select name="conceito">
          <option value="A">A – Aprovação com pequenas reformulações</option>
          <option value="B">B – Aprovação com reformulações estruturais</option>
          <option value="C">C – Aprovação com reformulações estruturais e metodológicas</option>
          <option value="D">D – Reprovação</option>
        </select></div>
      <div class="f"><label>Resolução</label>
        <input type="text" name="{prefix}_resolucao"></div>
    </div>'''

def make_resultado(prefix):
    return f'''
    <div class="row r2">
      <div class="f"><label>Resultado *</label>
        <select name="resultado">
          <option value="APROVADO(A)">APROVADO(A)</option>
          <option value="REPROVADO(A)">REPROVADO(A)</option>
        </select></div>
      <div class="f"><label>Resolução</label>
        <input type="text" name="{prefix}_resolucao"></div>
    </div>'''

@app.route('/')
def index():
    html = HTML
    # inject forms
    for prefix, trab in [('qm','dissertação'),('qd','tese'),('dm','dissertação'),('dd','tese')]:
        is_qual = prefix in ('qm','qd')
        form_html = make_form(prefix, trab)
        if is_qual:
            form_html += make_conceito(prefix)
        else:
            form_html += make_resultado(prefix)
        html = html.replace(f'{{{{ form_base(\'{prefix}\', \'{trab}\') }}}}', form_html)
        html = html.replace(f'{{{{ conceito_field(\'{prefix}\') }}}}', '')
        html = html.replace(f'{{{{ resultado_field(\'{prefix}\') }}}}', '')
    return html

@app.route('/logo1')
def logo1(): return send_file(LOGO1, mimetype='image/png')

@app.route('/logo2')
def logo2(): return send_file(LOGO2, mimetype='image/png')

@app.route('/logo3')
def logo3(): return send_file(os.path.join(os.path.dirname(__file__), 'logo3.jpeg'), mimetype='image/jpeg')

@app.route('/gerar', methods=['POST'])
def gerar():
    form = request.form
    tipo  = form.get('tipo','defesa')
    nivel = form.get('nivel','mestrado')
    fmt   = form.get('fmt','pdf')
    prefix = {'qualificacao':{'mestrado':'qm','doutorado':'qd'},
              'defesa':       {'mestrado':'dm','doutorado':'dd'}}[tipo][nivel]

    d = {
        'tipo': tipo, 'nivel': nivel,
        'dia':  form.get('dia',''), 'mes': form.get('mes',''), 'ano': form.get('ano',''),
        'hora': form.get('hora',''), 'modalidade': form.get('modalidade','de forma virtual'),
        'discente': form.get('discente',''), 'titulo': form.get('titulo',''),
        'or_trat': form.get(f'{prefix}_or_trat','Prof. Dr.'),
        'or_nome': form.get(f'{prefix}_or_nome',''),
        'resolucao': form.get(f'{prefix}_resolucao',''),
    }
    if tipo == 'qualificacao':
        d['conceito'] = form.get('conceito','A')
    else:
        d['resultado'] = form.get('resultado','APROVADO(A)')

    membros = parse_membros(form, prefix)
    titulo  = TITLES[(tipo, nivel)]

    if tipo == 'defesa':
        linhas = texto_defesa(d)
    else:
        linhas = texto_qualificacao(d)

    nome_arquivo = titulo.replace(' ','_').replace('/','_')

    if fmt == 'pdf':
        buf = gerar_pdf(d, membros, titulo, linhas)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True,
                         download_name=f'{nome_arquivo}.pdf')
    else:
        buf = gerar_docx(d, membros, titulo, linhas)
        return send_file(buf,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name=f'{nome_arquivo}.docx')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
